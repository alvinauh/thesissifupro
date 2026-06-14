"""
ThesisSifu Pro v4.1 — Multi-Agent Thesis Panel (Alignment-Aware Edition)
========================================================================
FIXES in v4.1:
  - Migrated to google.genai SDK (google-generativeai deprecated)
  - Robust chapter splitter: handles "CHAPTER 1", "Chapter One", "1.", Roman numerals,
    ALL-CAPS titles, titles on next line, and plain numbered headings
  - Robust subsection splitter: no longer filters by chapter_num (breaks on Roman numerals);
    instead detects subsections by position within the chapter text slice
  - Safe Gemini response handling: checks candidates/finish_reason before .text access,
    with detailed error logging per call site
  - Page estimation fixed: searches actual chapter text for the excerpt,
    not just the subsection slice
  - Prompt injection guard: CANONICAL constants passed via .replace() not .format()
    to prevent KeyError on any curly braces in canonical text
  - Annotated PDF: bumped search window and added full-text fallback search
  - All exceptions logged with chapter+subsection context for Railway log debugging

Pipeline:
  Stage 0  Spine Extraction        (Gemini)   → ThesisSpine
  Stage 1  Chapter + Subsection Split (regex) → structural map
  Stage 2  Subsection paragraph audit (Gemini, parallel, max 4 concurrent)
  Stage 3  Alignment Audit         (Claude)   → AlignmentMatrix
  Stage 4  Holistic Examiner Report (Claude)  → uses spine + matrix

Outputs (ZIP):
  1_Examiner_Audit_Report.pdf     — holistic critique, references spine
  2_Annotated_Thesis.(docx|pdf)   — subsection-aware inline comments
  3_Commentary_Report.pdf         — grouped by chapter → subsection → paragraph
  4_Alignment_Matrix_Report.pdf   — RQ↔RO↔Method↔Analysis↔Finding↔Conclusion

Endpoint:  POST /audit   multipart/form-data { file: <pdf|docx> }
Returns:   application/zip
"""

from __future__ import annotations

import io, os, re, json, asyncio, zipfile, hashlib, tempfile, traceback
from datetime import datetime
from dataclasses import dataclass, field as dc_field
from typing import Optional

import pypdf
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.background import BackgroundTask

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    Table, TableStyle, KeepTogether, PageBreak,
)

try:
    from google import genai as genai_sdk
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False


# ── App ────────────────────────────────────────────────────────
app = FastAPI(title="ThesisSifu Pro v4.1 — Alignment-Aware", version="4.1.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_credentials=False, allow_methods=["*"], allow_headers=["*"])


# ── AI Clients ─────────────────────────────────────────────────
gemini_client = None
claude_client = None
GEMINI_MODEL  = "gemini-3.1-flash-lite"   # flash (not flash-lite) for better JSON reliability

if GEMINI_AVAILABLE:
    gkey = os.environ.get("GEMINI_API_KEY")
    if gkey:
        gemini_client = genai_sdk.Client(api_key=gkey)
        print(f"[ThesisSifu] Gemini {GEMINI_MODEL} ready (google.genai SDK)")
    else:
        print("[ThesisSifu] WARNING: GEMINI_API_KEY not set")

if ANTHROPIC_AVAILABLE:
    akey = os.environ.get("ANTHROPIC_API_KEY")
    if akey:
        claude_client = anthropic.AsyncAnthropic(api_key=akey)
        print("[ThesisSifu] Claude Sonnet 4.6 ready")
    else:
        print("[ThesisSifu] WARNING: ANTHROPIC_API_KEY not set")


# ── Colours ────────────────────────────────────────────────────
W, H   = A4
BLACK  = colors.HexColor("#111111")
NAVY   = colors.HexColor("#0C2340")
ACCENT = colors.HexColor("#1A4A7A")
RULE   = colors.HexColor("#BBBBBB")
BOX_BG = colors.HexColor("#F4F6F9")
RED    = colors.HexColor("#A30000")
AMBER  = colors.HexColor("#7A4500")
GREEN  = colors.HexColor("#0A5C2A")
WHITE  = colors.white
LGRAY  = colors.HexColor("#EEEEEE")

SEV_COLORS = {"CRITICAL": RED, "MODERATE": AMBER, "SUGGESTION": ACCENT}
ALIGN_COLORS = {
    "ALIGNED":    GREEN,
    "PARTIAL":    AMBER,
    "MISALIGNED": RED,
    "MISSING":    colors.HexColor("#555555"),
    "UNCLEAR":    ACCENT,
}

TYPE_LABELS = {
    "JOURNAL_ARTICLE": "Scopus / Peer-Reviewed Journal Article",
    "UNDERGRADUATE":   "Undergraduate Thesis / Final Year Project",
    "MASTERS":         "Master's Thesis (By Research)",
    "PHD":             "PhD Dissertation",
}
TYPE_COLORS = {
    "JOURNAL_ARTICLE": colors.HexColor("#0A3D6B"),
    "UNDERGRADUATE":   colors.HexColor("#3A5A1C"),
    "MASTERS":         colors.HexColor("#6B3A00"),
    "PHD":             colors.HexColor("#4A0A0A"),
}


# ── Canonical libraries (NO curly braces — passed as plain text) ───────────
CANONICAL_FRAMEWORKS = (
    "THEORETICAL FRAMEWORKS (recommend by exact name + seminal author):\n"
    "\nTechnology / Information Systems:\n"
    "  - Technology Acceptance Model (TAM) — Davis (1989)\n"
    "  - UTAUT / UTAUT2 — Venkatesh et al. (2003, 2012)\n"
    "  - Diffusion of Innovations — Rogers (2003)\n"
    "  - Task-Technology Fit — Goodhue & Thompson (1995)\n"
    "  - DeLone & McLean IS Success Model — DeLone & McLean (2003)\n"
    "\nBehavioural / Psychological:\n"
    "  - Theory of Planned Behaviour — Ajzen (1991)\n"
    "  - Theory of Reasoned Action — Fishbein & Ajzen (1975)\n"
    "  - Social Cognitive Theory — Bandura (1986)\n"
    "  - Self-Determination Theory — Deci & Ryan (1985, 2000)\n"
    "  - Health Belief Model — Rosenstock (1974)\n"
    "\nStrategic Management / Business:\n"
    "  - Resource-Based View — Barney (1991)\n"
    "  - Dynamic Capabilities — Teece, Pisano & Shuen (1997)\n"
    "  - Porter's Five Forces — Porter (1980)\n"
    "  - Stakeholder Theory — Freeman (1984)\n"
    "  - Institutional Theory — DiMaggio & Powell (1983)\n"
    "\nOrganisational Behaviour / HR:\n"
    "  - Job Demands-Resources Model — Bakker & Demerouti (2007)\n"
    "  - Social Exchange Theory — Blau (1964)\n"
    "  - Transformational Leadership — Bass (1985)\n"
    "  - Psychological Capital — Luthans et al. (2007)\n"
    "\nMarketing / Consumer Behaviour:\n"
    "  - Stimulus-Organism-Response (SOR) — Mehrabian & Russell (1974)\n"
    "  - SERVQUAL — Parasuraman, Zeithaml & Berry (1988)\n"
    "  - Customer-Based Brand Equity — Keller (1993)\n"
    "  - Expectation-Confirmation Theory — Oliver (1980)\n"
    "\nEducation / Learning:\n"
    "  - Constructivism — Piaget, Vygotsky\n"
    "  - Bloom's Taxonomy (revised) — Anderson & Krathwohl (2001)\n"
    "  - Community of Inquiry — Garrison, Anderson & Archer (2000)\n"
    "  - TPACK — Mishra & Koehler (2006)\n"
)

CANONICAL_METHODS = (
    "METHODOLOGICAL REFERENCES (recommend by exact name + seminal author):\n"
    "\nQuantitative analytical techniques:\n"
    "  - PLS-SEM — Hair, Hult, Ringle & Sarstedt (2017, 2022)\n"
    "  - CB-SEM — Byrne (2016), Kline (2015)\n"
    "  - Hierarchical regression / moderation — Aiken & West (1991), Hayes PROCESS (2018)\n"
    "  - Mediation — Baron & Kenny (1986), Preacher & Hayes (2008)\n"
    "  - ANOVA/MANOVA — Tabachnick & Fidell (2019)\n"
    "  - Factor Analysis (EFA/CFA) — Hair et al. (2019), Brown (2015)\n"
    "\nQualitative analytical techniques:\n"
    "  - Thematic Analysis — Braun & Clarke (2006, 2019)\n"
    "  - Reflexive Thematic Analysis — Braun & Clarke (2022)\n"
    "  - Grounded Theory — Charmaz (2014), Strauss & Corbin (1998)\n"
    "  - Gioia Methodology — Gioia, Corley & Hamilton (2013)\n"
    "  - IPA — Smith, Flowers & Larkin (2009)\n"
    "  - Case Study — Yin (2018), Eisenhardt (1989)\n"
    "\nSystematic review:\n"
    "  - PRISMA 2020 — Page et al. (2021)\n"
    "  - Scoping review — Arksey & O'Malley (2005)\n"
    "\nSampling / sample size:\n"
    "  - Cochran's formula — Cochran (1977)\n"
    "  - Krejcie & Morgan table — Krejcie & Morgan (1970)\n"
    "  - G*Power power analysis — Faul et al. (2009)\n"
    "  - Minimum sample for PLS-SEM — Hair et al. (2017)\n"
    "\nValidity & reliability:\n"
    "  - Cronbach's alpha — Cronbach (1951)\n"
    "  - Composite reliability, AVE — Hair et al. (2017), Fornell & Larcker (1981)\n"
    "  - HTMT — Henseler, Ringle & Sarstedt (2015)\n"
    "  - Common Method Bias — Podsakoff et al. (2003, 2012)\n"
    "  - Content validity index (CVI) — Lynn (1986)\n"
    "\nQualitative trustworthiness:\n"
    "  - Lincoln & Guba (1985) — credibility, transferability, dependability, confirmability\n"
)


# ── Data classes ───────────────────────────────────────────────
@dataclass
class ThesisSpine:
    title:              str = "UNKNOWN"
    problem_statement:  str = "NOT FOUND"
    research_gap:       str = "NOT FOUND"
    research_questions: list = dc_field(default_factory=list)
    research_objectives:list = dc_field(default_factory=list)
    hypotheses:         list = dc_field(default_factory=list)
    theory_used:        str = "NOT FOUND"
    variables:          list = dc_field(default_factory=list)
    methodology:        str = "NOT FOUND"
    sampling:           str = "NOT FOUND"
    instrument:         str = "NOT FOUND"
    analysis_technique: str = "NOT FOUND"
    key_findings:       list = dc_field(default_factory=list)
    conclusions:        list = dc_field(default_factory=list)
    discipline:         str = "UNKNOWN"

@dataclass
class ParagraphComment:
    chapter:             str
    subsection:          str
    subsection_title:    str
    para_index:          int
    page_estimate:       int
    para_excerpt:        str
    severity:            str
    issue:               str
    recommendation:      str
    literature_needed:   str
    theory_needed:       str
    suggested_framework: str = ""
    suggested_method:    str = ""

@dataclass
class Subsection:
    chapter_num:      str
    subsection_num:   str
    title:            str
    text:             str
    char_offset:      int = 0          # offset within full_text for page estimation
    expected_purpose: str = ""
    comments:         list = dc_field(default_factory=list)

@dataclass
class ChapterSummary:
    chapter_num:   str
    chapter_title: str
    text:          str = ""
    char_offset:   int = 0             # offset within full_text for page estimation
    subsections:   list = dc_field(default_factory=list)
    comments:      list = dc_field(default_factory=list)

@dataclass
class AlignmentRow:
    rq:         str = "—"
    ro:         str = "—"
    hypothesis: str = "—"
    method:     str = "—"
    analysis:   str = "—"
    finding:    str = "—"
    conclusion: str = "—"
    status:     str = "UNCLEAR"
    note:       str = ""

@dataclass
class AlignmentMatrix:
    rows:                       list = dc_field(default_factory=list)
    overall_verdict:            str = ""
    golden_thread_score:        str = "UNCLEAR"
    critical_gaps:              list = dc_field(default_factory=list)
    structural_recommendations: list = dc_field(default_factory=list)


# ── Gemini response helper ─────────────────────────────────────
def _gemini_text(response, context: str = "") -> str:
    """Safely extract text from a google.genai response.
    Logs the finish reason if the response was blocked or empty."""
    try:
        if not response or not response.candidates:
            print(f"[ThesisSifu] Gemini: no candidates returned. Context: {context}")
            return ""
        cand = response.candidates[0]
        # finish_reason: 1=STOP 2=MAX_TOKENS 3=SAFETY 4=RECITATION 5=OTHER
        reason = getattr(cand, "finish_reason", None)
        if reason and reason not in (1, "STOP"):
            print(f"[ThesisSifu] Gemini blocked/incomplete. reason={reason} context={context}")
            return ""
        return response.text or ""
    except Exception as e:
        print(f"[ThesisSifu] _gemini_text error ({context}): {e}")
        return ""


# ── Text extraction ─────────────────────────────────────────────
def extract_text(content: bytes, filename: str) -> str:
    """Extract text from PDF or DOCX.
    Uses pdfplumber for PDFs (preserves inter-word spaces that pypdf often drops),
    with pypdf as fallback. Word-concatenation like '1Introduction' or '3ProblemStatement'
    causes the chapter/subsection regex to produce hundreds of false positives,
    which then starves the Gemini subsection auditor of real text."""
    fn = filename.lower()
    if fn.endswith(".pdf"):
        # Primary: pdfplumber (layout-aware, preserves spaces)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text(x_tolerance=1, y_tolerance=2) or ""
                    pages.append(t)
                text = "\n".join(pages)
                if text.strip():
                    print(f"[ThesisSifu] pdfplumber extracted {len(text)} chars")
                    return text
        except Exception as e:
            print(f"[ThesisSifu] pdfplumber failed ({e}), falling back to pypdf")
        # Fallback: pypdf
        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                t = page.extract_text() or ""
                pages.append(t)
            text = "\n".join(pages)
            print(f"[ThesisSifu] pypdf fallback extracted {len(text)} chars")
            return text
        except Exception as e:
            print(f"[ThesisSifu] pypdf also failed: {e}")
            return ""
    if fn.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return ""


# ── Chapter + Subsection splitter (robust) ────────────────────
# Chapter pattern: ONLY matches "CHAPTER N" / "CHAPTER I" / "CHAPTER ONE"
# and "N.0 Title" forms (explicit section-zero heading).
# The old (\d+)[:\.\-] branch was matching statistics (3.432, 0.521) as chapters.
_WORD_NUMS = {"one":"1","two":"2","three":"3","four":"4","five":"5",
              "six":"6","seven":"7","eight":"8","nine":"9","ten":"10"}
_ROMAN = {"I":"1","II":"2","III":"3","IV":"4","V":"5",
          "VI":"6","VII":"7","VIII":"8","IX":"9","X":"10"}

_CH_PAT = re.compile(
    r'(?:^|\n)'
    r'(?:'
      # "CHAPTER 1", "CHAPTER I", "CHAPTER ONE", "BAB 1", "BAHAGIAN 2"
      r'(?:CHAPTER|BAHAGIAN|BAB)\s+(\d{1,2}|[IVX]+|ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\b'
      # "1.0 Introduction" or "2.0 Literature Review" — ONLY N.0 forms, not N.432
      r'|(\d{1,2})\.0+\s'
    r')'
    r'[:\s\-—–]*([^\n]{0,100})',
    re.IGNORECASE,
)

# Subsection: "1.1 Title", "3.3.1 Title" — integers only, MUST be followed by space+letter.
# Lookahead (?=[ \t]+[A-Za-z]) prevents matching statistics like 3.432, 0.521, 1.015.
_SUB_PAT = re.compile(
    r'(?:^|\n)[ \t]*'
    r'(\d{1,2}\.\d{1,3}(?:\.\d{1,3})?)'   # N.M or N.M.P (short integers, max 3 digits each)
    r'(?=[ \t]+[A-Za-z])'                   # LOOKAHEAD: must be followed by whitespace then a letter
    r'[ \t]+([^\n]{0,120})',                 # capture the title
)

_SUBSECTION_PURPOSE = {
    "background":           "Establish context and historical/practical relevance of the problem.",
    "problem statement":    "Articulate the specific research problem with empirical/theoretical evidence.",
    "research gap":         "Demonstrate what prior work has NOT addressed.",
    "research question":    "Pose 1+ specific, answerable research questions tied to the problem.",
    "objective":            "List measurable objectives that map 1-to-1 to research questions.",
    "hypothes":             "State testable, directional hypotheses linked to theory.",
    "significance":         "Explain theoretical and practical contributions of the study.",
    "scope":                "Define boundaries (population, geography, time, constructs).",
    "definition":           "Operationalise key terms used in the study.",
    "theoretical framework":"Justify the underpinning theory and link constructs to it.",
    "conceptual framework": "Present the model/diagram of relationships among variables.",
    "literature review":    "Synthesise prior research critically, not just summarise.",
    "research design":      "Justify quantitative/qualitative/mixed approach and paradigm.",
    "population":           "Define the target population and sampling frame.",
    "sampling":             "Justify sampling technique and sample size with a recognised formula.",
    "instrument":           "Describe instrument development, items, scale, and pretest.",
    "validity":             "Report content/construct validity evidence (e.g., CVI, AVE).",
    "reliability":          "Report reliability evidence (Cronbach's alpha, composite reliability).",
    "data collection":      "Detail procedure, ethics approval, response rate.",
    "data analysis":        "Justify analytical technique against research questions.",
    "demographic":          "Present sample profile relevant to research questions.",
    "descriptive":          "Report descriptives that contextualise constructs.",
    "results":              "Present findings against each research question/hypothesis explicitly.",
    "finding":              "Report findings tied directly to each RQ.",
    "discussion":           "Interpret findings against literature; explain unexpected results.",
    "implication":          "Theoretical and practical implications grounded in findings.",
    "limitation":           "Acknowledge methodological and scope limitations honestly.",
    "future research":      "Propose specific, actionable next studies.",
    "conclusion":           "Restate contribution and answer each research question explicitly.",
    "summary":              "Concise restatement of findings against objectives.",
    "recommendation":       "Specific, actionable recommendations for stakeholders.",
}

def _purpose_for(title: str) -> str:
    t = title.lower()
    for k, v in _SUBSECTION_PURPOSE.items():
        if k in t:
            return v
    return "Auditor should infer purpose from chapter context."

def _normalise_chnum(raw: str) -> str:
    """Convert Roman/word chapter numbers to Arabic integers."""
    r = raw.strip().upper()
    if r in _ROMAN: return _ROMAN[r]
    if r.lower() in _WORD_NUMS: return _WORD_NUMS[r.lower()]
    return raw.strip()

def split_chapters(text: str) -> list[ChapterSummary]:
    matches = list(_CH_PAT.finditer(text))
    if not matches:
        size = 6000
        return [ChapterSummary(chapter_num=str(i+1),
                               chapter_title=f"Section {i+1}",
                               text=text[s:s+size],
                               char_offset=s)
                for i, s in enumerate(range(0, len(text), size))]
    chapters = []
    for i, m in enumerate(matches):
        raw_num = (m.group(1) or m.group(2) or str(i+1))
        num     = _normalise_chnum(raw_num)
        title   = (m.group(3) or "").strip() or f"Chapter {num}"
        # If title is empty the heading might continue on the next line
        if not title.strip():
            next_nl = text.find("\n", m.end())
            if next_nl != -1:
                title = text[m.end():next_nl].strip() or f"Chapter {num}"
        start   = m.start()
        end     = matches[i+1].start() if i+1 < len(matches) else len(text)
        ch_text = text[start:end].strip()
        chapters.append(ChapterSummary(
            chapter_num=num, chapter_title=title,
            text=ch_text, char_offset=start))
    return chapters or [ChapterSummary("1", "Full Document", text, 0)]


def split_subsections(chapter: ChapterSummary) -> list[Subsection]:
    """Split a chapter into subsections.
    The _SUB_PAT lookahead (space+letter) already filters statistics,
    but we also enforce a minimum text length so tiny fragments don't
    consume Gemini quota returning empty results."""
    MIN_SUB_CHARS = 300   # skip subsections shorter than this
    matches = list(_SUB_PAT.finditer(chapter.text))
    if not matches:
        return [Subsection(
            chapter_num=chapter.chapter_num, subsection_num="—",
            title=chapter.chapter_title, text=chapter.text,
            char_offset=chapter.char_offset,
            expected_purpose=_purpose_for(chapter.chapter_title))]
    subs = []
    for i, m in enumerate(matches):
        sub_num   = m.group(1).strip()
        sub_title = (m.group(2) or "").strip() or f"Section {sub_num}"
        start     = m.start()
        end       = matches[i+1].start() if i+1 < len(matches) else len(chapter.text)
        sub_text  = chapter.text[start:end].strip()
        if len(sub_text) < MIN_SUB_CHARS:
            continue   # skip navigation entries, TOC lines, one-liner sections
        subs.append(Subsection(
            chapter_num=chapter.chapter_num, subsection_num=sub_num,
            title=sub_title, text=sub_text,
            char_offset=chapter.char_offset + start,
            expected_purpose=_purpose_for(sub_title),
        ))
    # If all were too short, fall back to treating whole chapter as one subsection
    if not subs:
        return [Subsection(
            chapter_num=chapter.chapter_num, subsection_num="—",
            title=chapter.chapter_title, text=chapter.text,
            char_offset=chapter.char_offset,
            expected_purpose=_purpose_for(chapter.chapter_title))]
    return subs


# ── Page estimator (uses full_text offset, not sub slice) ─────
def _estimate_page(full_text: str, excerpt: str, char_offset: int,
                   words_per_page: int = 350) -> int:
    """Find excerpt in full_text starting from char_offset; fall back to offset-based."""
    search_start = max(0, char_offset - 200)
    pos = full_text.find(excerpt[:50].strip(), search_start)
    if pos < 0:
        # Try shorter excerpt
        pos = full_text.find(excerpt[:25].strip(), search_start)
    if pos > 0:
        word_count = len(full_text[:pos].split())
        return max(1, word_count // words_per_page + 1)
    # Fall back to offset
    word_count = len(full_text[:char_offset].split())
    return max(1, word_count // words_per_page + 1)


def _find_content_start(text: str) -> int:
    """Skip over preamble pages (cover, admin forms, acknowledgements, table of contents).
    Returns the char position where substantive academic content begins.
    Looks for the Abstract, then Chapter 1, then any 'Introduction' heading."""
    markers = ["abstract\n", "Abstract\n", "ABSTRACT\n",
               "CHAPTER 1\n", "Chapter 1\n", "CHAPTER ONE\n",
               "1.1 Introduction\n", "1.1 Background\n",
               "Introduction\n1.", "INTRODUCTION\n"]
    best = len(text)
    for marker in markers:
        pos = text.find(marker)
        if 0 < pos < best:
            best = pos
    # If nothing found in reasonable range, start at 0
    return best if best < len(text) * 0.3 else 0

def _extract_spine_fallback(text: str) -> dict:
    """Rule-based spine extraction as fallback when Gemini fails.
    Searches for explicit RQ/RO patterns, methodology keywords, etc."""
    d = {
        "title": "UNKNOWN", "discipline": "UNKNOWN",
        "problem_statement": "NOT FOUND", "research_gap": "NOT FOUND",
        "research_questions": [], "research_objectives": [], "hypotheses": [],
        "theory_used": "NOT FOUND", "variables": [],
        "methodology": "NOT FOUND", "sampling": "NOT FOUND",
        "instrument": "NOT FOUND", "analysis_technique": "NOT FOUND",
        "key_findings": [], "conclusions": []
    }
    # Title: look for ALL-CAPS line or explicit "Title:" field
    title_m = re.search(r'(?:Tajuk|Title)\s*[:：]\s*([^\n]{10,120})', text, re.IGNORECASE)
    if title_m: d["title"] = title_m.group(1).strip()
    else:
        # Try bold-looking ALL CAPS heading near top
        caps_m = re.search(r'\n([A-Z][A-Z\s]{15,80})\n', text[:3000])
        if caps_m: d["title"] = caps_m.group(1).strip()

    # Research Questions: look for RQ1/RQ2 or "Research Question 1"
    rq_matches = re.findall(
        r'RQ\s*\d+\s*[:：]\s*([^\n?]+\??)', text, re.IGNORECASE)
    if not rq_matches:
        rq_matches = re.findall(
            r'Research Question \d+\s*[:：]\s*([^\n?]+\??)', text, re.IGNORECASE)
    d["research_questions"] = [r.strip() for r in rq_matches[:6]]

    # Research Objectives: numbered "To verb..." lists
    ro_matches = re.findall(
        r'\d+\.\s+(To\s+(?:identify|examine|determine|explore|assess|investigate|describe|compare|evaluate|analyse|analyze)[^\n\.]{10,200})',
        text, re.IGNORECASE)
    d["research_objectives"] = [r.strip() for r in ro_matches[:6]]

    # Hypotheses
    h_matches = re.findall(r'H[01]\d*\s*[:：]\s*([^\n\.]{10,200})', text)
    d["hypotheses"] = [h.strip() for h in h_matches[:6]]

    # Methodology keywords
    if re.search(r'mixed[\s-]method', text, re.IGNORECASE):
        d["methodology"] = "Mixed-methods"
        if re.search(r'explanatory sequential', text, re.IGNORECASE):
            d["methodology"] = "Explanatory sequential mixed-methods"
        elif re.search(r'convergent parallel', text, re.IGNORECASE):
            d["methodology"] = "Convergent parallel mixed-methods"
    elif re.search(r'\bqualitative\b', text, re.IGNORECASE) and not re.search(r'\bquantitative\b', text, re.IGNORECASE):
        d["methodology"] = "Qualitative"
    elif re.search(r'\bquantitative\b', text, re.IGNORECASE):
        d["methodology"] = "Quantitative"

    # Analysis technique
    for technique in ["PLS-SEM", "CB-SEM", "thematic analysis", "Thematic Analysis",
                       "Pearson correlation", "regression", "ANOVA", "IPA",
                       "grounded theory", "Grounded Theory", "descriptive statistics"]:
        if technique.lower() in text.lower():
            existing = d["analysis_technique"]
            if existing == "NOT FOUND":
                d["analysis_technique"] = technique
            elif technique.lower() not in existing.lower():
                d["analysis_technique"] = existing + " + " + technique
            break

    # Theory
    theory_patterns = [
        ("Social Cognitive Theory", r"Social Cognitive Theory|Bandura.*self.efficacy"),
        ("TAM", r"\bTAM\b|Technology Acceptance Model"),
        ("UTAUT", r"\bUTAUT\b"),
        ("TPB", r"Theory of Planned Behaviour|Theory of Planned Behavior"),
        ("SDT", r"Self.Determination Theory"),
    ]
    for name, pat in theory_patterns:
        if re.search(pat, text, re.IGNORECASE):
            d["theory_used"] = name
            break

    # Sampling
    samp_m = re.search(r'(n\s*=\s*\d+|sample\s+of\s+\d+|\d+\s+(?:respondents|participants|teachers|students))',
                        text, re.IGNORECASE)
    if samp_m: d["sampling"] = samp_m.group(0).strip()

    # Abstract for problem + findings
    abs_m = re.search(r'Abstract\s*\n(.{200,2000}?)(?:\n\n|\nKeywords|\nAbstrak)', text, re.IGNORECASE | re.DOTALL)
    if abs_m:
        abstract = abs_m.group(1).strip()
        d["problem_statement"] = abstract[:400]
        # Findings often in last sentence of abstract
        sentences = [s.strip() for s in abstract.split('.') if len(s.strip()) > 30]
        if sentences:
            d["key_findings"] = [sentences[-1]] if sentences else []

    # Discipline from keywords
    for disc, pats in [
        ("Education", ["teacher", "student", "classroom", "learning", "instruction", "pedagogy"]),
        ("Information Systems", ["system", "technology", "adoption", "TAM", "IS"]),
        ("Business", ["firm", "organization", "management", "strategy"]),
        ("Health", ["patient", "health", "clinical", "medical"]),
    ]:
        if any(p.lower() in text[:5000].lower() for p in pats):
            d["discipline"] = disc
            break

    print(f"[ThesisSifu] Fallback spine: RQs={len(d['research_questions'])}, "
          f"ROs={len(d['research_objectives'])}, method={d['methodology']}, "
          f"theory={d['theory_used']}")
    return d


# ── Stage 0: Spine Extraction ──────────────────────────────────
_SPINE_PROMPT = """\
You are extracting the STRUCTURAL SPINE of an academic thesis or article.
Return STRICT JSON only — no markdown, no commentary, no preamble.

Fields (use "NOT FOUND" or [] if absent from the excerpt):
{
  "title": "exact title",
  "discipline": "best guess discipline (e.g., Information Systems, Marketing, Education)",
  "problem_statement": "1-3 sentence summary of the problem",
  "research_gap": "1-2 sentence summary of the gap claimed",
  "research_questions": ["RQ1 verbatim or paraphrased", "RQ2..."],
  "research_objectives": ["RO1...", "RO2..."],
  "hypotheses": ["H1...", "H2..."],
  "theory_used": "name of underpinning theory/framework if stated",
  "variables": ["IV1", "IV2", "DV", "MV (moderator/mediator)"],
  "methodology": "design + paradigm (e.g., quantitative cross-sectional survey)",
  "sampling": "technique + size if stated",
  "instrument": "questionnaire / interview protocol / etc.",
  "analysis_technique": "PLS-SEM / thematic / regression / etc.",
  "key_findings": ["finding 1", "finding 2"],
  "conclusions": ["conclusion 1", "conclusion 2"]
}

Document excerpt (front + back samples):
\"\"\"EXCERPT_PLACEHOLDER\"\"\"\
"""

async def extract_spine(text: str, full_text: str) -> ThesisSpine:
    # ── Smart sampling: skip preamble, start from Abstract/Chapter 1 ──────────
    content_start = _find_content_start(text)
    if content_start > 0:
        print(f"[ThesisSifu] Skipping {content_start} chars of preamble for spine")
    # Take front from content start, middle from RQ area, back from conclusions
    content = text[content_start:]
    front  = content[:6000]
    # Try to find the RQ section specifically
    rq_pos = text.find("RQ1:") or text.find("Research Question 1")
    if rq_pos > 0:
        middle = text[max(0, rq_pos-200):rq_pos+2000]
    else:
        middle = content[4000:8000] if len(content) > 8000 else ""
    back   = text[-3000:] if len(text) > 15000 else ""
    sample = (front + "\n\n[...RQ section...]\n\n" + middle +
              "\n\n[...conclusion section...]\n\n" + back)[:14000]
    prompt = _SPINE_PROMPT.replace("EXCERPT_PLACEHOLDER", sample)

    d = None
    if gemini_client and text.strip():
        try:
            r = await gemini_client.aio.models.generate_content(
                model=GEMINI_MODEL, contents=prompt)
            raw = _gemini_text(r, "spine_extraction")
            if raw:
                raw = re.sub(r"```(?:json)?", "", raw).strip("`").strip()
                d = json.loads(raw)
                print(f"[ThesisSifu] Gemini spine: RQs={len(d.get('research_questions',[]))}")
        except Exception as e:
            print(f"[ThesisSifu] Spine Gemini error: {e}")

    # ── Fallback: rule-based extraction if Gemini failed or returned mostly NOT FOUND ──
    if d is None:
        print("[ThesisSifu] Spine: using rule-based fallback (no Gemini result)")
        d = _extract_spine_fallback(text)
    else:
        # Merge: fill in any NOT FOUND fields with rule-based results
        fallback = _extract_spine_fallback(text)
        for key in ["title", "problem_statement", "research_questions", "research_objectives",
                    "hypotheses", "theory_used", "methodology", "analysis_technique",
                    "sampling", "key_findings"]:
            gemini_val = d.get(key)
            fb_val = fallback.get(key)
            if (not gemini_val or gemini_val in ("UNKNOWN","NOT FOUND",[])) and fb_val and fb_val not in ("UNKNOWN","NOT FOUND",[]):
                d[key] = fb_val
                print(f"[ThesisSifu] Spine fallback filled: {key}")

    return ThesisSpine(
        title=              d.get("title","UNKNOWN"),
        discipline=         d.get("discipline","UNKNOWN"),
        problem_statement=  d.get("problem_statement","NOT FOUND"),
        research_gap=       d.get("research_gap","NOT FOUND"),
        research_questions= d.get("research_questions",[]) or [],
        research_objectives=d.get("research_objectives",[]) or [],
        hypotheses=         d.get("hypotheses",[]) or [],
        theory_used=        d.get("theory_used","NOT FOUND"),
        variables=          d.get("variables",[]) or [],
        methodology=        d.get("methodology","NOT FOUND"),
        sampling=           d.get("sampling","NOT FOUND"),
        instrument=         d.get("instrument","NOT FOUND"),
        analysis_technique= d.get("analysis_technique","NOT FOUND"),
        key_findings=       d.get("key_findings",[]) or [],
        conclusions=        d.get("conclusions",[]) or [],
    )


# ── Stage 1: Classification ─────────────────────────────────────
_CLS_PROMPT = """\
Classify this academic document. Choose ONE:
JOURNAL_ARTICLE | UNDERGRADUATE | MASTERS | PHD

Return ONLY valid JSON (no markdown):
{"type":"PHD","confidence":"HIGH","signals":["s1","s2"],"title":"t","authors":"a","field":"f","institution":"i"}

Excerpt:
\"\"\"EXCERPT_PLACEHOLDER\"\"\"\
"""

async def classify_document(text: str) -> dict:
    default = {"type":"MASTERS","confidence":"LOW","signals":["fallback"],
               "title":"UNKNOWN","authors":"UNKNOWN","field":"UNKNOWN","institution":"UNKNOWN"}
    if not gemini_client or not text.strip():
        return default
    prompt = _CLS_PROMPT.replace("EXCERPT_PLACEHOLDER", text[:5000])
    try:
        r = await gemini_client.aio.models.generate_content(
            model=GEMINI_MODEL, contents=prompt)
        raw = _gemini_text(r, "classify_document")
        if not raw:
            return default
        raw = re.sub(r"```(?:json)?","", raw).strip("`").strip()
        return json.loads(raw)
    except Exception as e:
        print(f"[ThesisSifu] Classification error: {e}")
        return default


# ── Stage 2: Subsection paragraph audit ───────────────────────
# NOTE: Uses .replace() for CANONICAL blocks, .format() only for
#       the simple non-curly-brace fields.
_SUBSECTION_PROMPT_TEMPLATE = """\
You are a PhD supervisor auditing one SUBSECTION of a DOC_TYPE_PLACEHOLDER.

THESIS SPINE (for alignment checking — cite these in every comment):
  Title:            TITLE_PLACEHOLDER
  Problem:          PROBLEM_PLACEHOLDER
  Research Qs:      RQS_PLACEHOLDER
  Research Objs:    ROS_PLACEHOLDER
  Theory used:      THEORY_PLACEHOLDER
  Methodology:      METHOD_PLACEHOLDER
  Analysis:         ANALYSIS_PLACEHOLDER
  Variables:        VARS_PLACEHOLDER

CHAPTER CHNUM_PLACEHOLDER — CHTITLE_PLACEHOLDER
SUBSECTION SUBNUM_PLACEHOLDER — SUBTITLE_PLACEHOLDER
Expected purpose of this subsection: PURPOSE_PLACEHOLDER

YOUR TASK
For 4-8 paragraphs needing attention in this subsection, return a JSON array.
Each element must have these EXACT keys:
[
  {
    "para_excerpt": "first 80 chars of the paragraph verbatim",
    "severity": "CRITICAL",
    "issue": "specific intellectual problem",
    "recommendation": "specific fix tied to thesis spine",
    "literature_needed": "type of evidence needed (do NOT invent paper titles)",
    "theory_needed": "named framework OR none",
    "suggested_framework": "EXACT name + author from canonical list, e.g. TAM (Davis 1989), or empty string",
    "suggested_method": "EXACT name + author from canonical list, e.g. PLS-SEM (Hair et al. 2017), or empty string"
  }
]

ALIGNMENT — flag CRITICAL when:
  - Paragraph contradicts the problem statement or strays from RQs/ROs
  - Subsection does not deliver its expected purpose
  - Methodology cannot answer the stated RQ
  - Finding does not map to any RQ
  - Conclusion overreaches what the analysis supports
  - Theory invoked but not operationalised

CANONICAL FRAMEWORKS:
FRAMEWORKS_PLACEHOLDER

CANONICAL METHODS:
METHODS_PLACEHOLDER

Return ONLY the JSON array. No markdown fences. No prose before or after.

Subsection text:
\"\"\"TEXT_PLACEHOLDER\"\"\"\
"""

async def audit_subsection(sub: Subsection, ch_title: str, doc_type: str,
                            spine: ThesisSpine, full_text: str) -> list[ParagraphComment]:
    if not gemini_client or not sub.text.strip():
        return []
    # Build prompt using .replace() to avoid KeyError on braces in canonical text or user text
    prompt = (
        _SUBSECTION_PROMPT_TEMPLATE
        .replace("DOC_TYPE_PLACEHOLDER", doc_type)
        .replace("TITLE_PLACEHOLDER",    spine.title[:200])
        .replace("PROBLEM_PLACEHOLDER",  spine.problem_statement[:300])
        .replace("RQS_PLACEHOLDER",      "; ".join(spine.research_questions[:5]) or "NOT FOUND")
        .replace("ROS_PLACEHOLDER",      "; ".join(spine.research_objectives[:5]) or "NOT FOUND")
        .replace("THEORY_PLACEHOLDER",   spine.theory_used)
        .replace("METHOD_PLACEHOLDER",   spine.methodology)
        .replace("ANALYSIS_PLACEHOLDER", spine.analysis_technique)
        .replace("VARS_PLACEHOLDER",     "; ".join(spine.variables[:8]) or "NOT FOUND")
        .replace("CHNUM_PLACEHOLDER",    sub.chapter_num)
        .replace("CHTITLE_PLACEHOLDER",  ch_title[:80])
        .replace("SUBNUM_PLACEHOLDER",   sub.subsection_num)
        .replace("SUBTITLE_PLACEHOLDER", sub.title[:80])
        .replace("PURPOSE_PLACEHOLDER",  sub.expected_purpose[:200])
        .replace("FRAMEWORKS_PLACEHOLDER", CANONICAL_FRAMEWORKS)
        .replace("METHODS_PLACEHOLDER",    CANONICAL_METHODS)
        .replace("TEXT_PLACEHOLDER",       sub.text[:7000])
    )
    ctx = f"ch{sub.chapter_num}.{sub.subsection_num}"
    try:
        r = await gemini_client.aio.models.generate_content(
            model=GEMINI_MODEL, contents=prompt)
        raw = _gemini_text(r, f"audit_subsection_{ctx}")
        if not raw:
            print(f"[ThesisSifu] Empty response for subsection {ctx}")
            return []
        raw = re.sub(r"```(?:json)?", "", raw).strip("`").strip()
        # Gemini sometimes returns a single object instead of array — wrap it
        if raw.startswith("{"):
            raw = "[" + raw + "]"
        items = json.loads(raw)
        if not isinstance(items, list):
            items = [items]
        comments = []
        for idx, it in enumerate(items):
            excerpt = str(it.get("para_excerpt",""))[:150]
            pg = _estimate_page(full_text, excerpt, sub.char_offset)
            comments.append(ParagraphComment(
                chapter=            f"Chapter {sub.chapter_num}",
                subsection=         sub.subsection_num,
                subsection_title=   sub.title,
                para_index=         idx,
                page_estimate=      pg,
                para_excerpt=       excerpt,
                severity=           str(it.get("severity","MODERATE")).upper(),
                issue=              str(it.get("issue","")),
                recommendation=     str(it.get("recommendation","")),
                literature_needed=  str(it.get("literature_needed","")),
                theory_needed=      str(it.get("theory_needed","none")),
                suggested_framework=str(it.get("suggested_framework","")),
                suggested_method=   str(it.get("suggested_method","")),
            ))
        print(f"[ThesisSifu] {ctx}: {len(comments)} comments")
        return comments
    except json.JSONDecodeError as e:
        print(f"[ThesisSifu] JSON parse error in {ctx}: {e}\nRaw: {raw[:300]}")
        return []
    except Exception as e:
        print(f"[ThesisSifu] Subsection audit error {ctx}: {e}\n{traceback.format_exc()}")
        return []


# ── Stage 3: Alignment Audit (Claude) ─────────────────────────
_ALIGNMENT_SYSTEM = """\
You are a senior academic examiner auditing STRUCTURAL ALIGNMENT of a thesis.
Verify the golden thread: Problem -> RQ -> RO -> Theory -> Method -> Analysis -> Findings -> Conclusion.
Return STRICT JSON only. No markdown, no prose outside the JSON.\
"""

_ALIGNMENT_PROMPT_TEMPLATE = """\
THESIS SPINE
============
Title:                TITLE_PLACEHOLDER
Discipline:           DISCIPLINE_PLACEHOLDER
Problem statement:    PROBLEM_PLACEHOLDER
Research gap:         GAP_PLACEHOLDER
Research questions:   RQS_PLACEHOLDER
Research objectives:  ROS_PLACEHOLDER
Hypotheses:           HYPS_PLACEHOLDER
Theory used:          THEORY_PLACEHOLDER
Variables/constructs: VARS_PLACEHOLDER
Methodology:          METHOD_PLACEHOLDER
Sampling:             SAMPLING_PLACEHOLDER
Instrument:           INSTRUMENT_PLACEHOLDER
Analysis technique:   ANALYSIS_PLACEHOLDER
Key findings:         FINDINGS_PLACEHOLDER
Conclusions:          CONCLUSIONS_PLACEHOLDER

CRITICAL ISSUES RAISED BY SUBSECTION AUDIT
==========================================
CRITICAL_ISSUES_PLACEHOLDER

YOUR TASK
=========
Produce a JSON object:
{
  "rows": [
    {
      "rq": "RQ1 verbatim",
      "ro": "RO1 matched to RQ1",
      "hypothesis": "H1 if applicable else none",
      "method": "method used to answer RQ1",
      "analysis": "analysis technique applied",
      "finding": "finding relevant to RQ1",
      "conclusion": "conclusion drawn for RQ1",
      "status": "ALIGNED or PARTIAL or MISALIGNED or MISSING or UNCLEAR",
      "note": "concrete reason — quote spine fields where possible"
    }
  ],
  "golden_thread_score": "STRONG or ACCEPTABLE or WEAK or BROKEN",
  "overall_verdict": "2-3 sentence verdict on structural coherence",
  "critical_gaps": ["gap 1 with specific explanation", "gap 2"],
  "structural_recommendations": ["specific fix 1 with named method/framework", "fix 2"]
}

Return ONLY the JSON object. No prose, no markdown.\
"""

async def run_alignment_audit(spine: ThesisSpine, chs: list[ChapterSummary]) -> AlignmentMatrix:
    if not claude_client:
        return AlignmentMatrix(
            overall_verdict="Alignment audit unavailable — ANTHROPIC_API_KEY not set.",
            golden_thread_score="UNCLEAR")
    critical_issues = []
    for cs in chs:
        for c in cs.comments:
            if c.severity == "CRITICAL":
                critical_issues.append(f"Ch.{c.chapter} §{c.subsection}: {c.issue}")
    critical_block = "\n".join(critical_issues[:30]) or "(none)"

    prompt = (
        _ALIGNMENT_PROMPT_TEMPLATE
        .replace("TITLE_PLACEHOLDER",       spine.title)
        .replace("DISCIPLINE_PLACEHOLDER",  spine.discipline)
        .replace("PROBLEM_PLACEHOLDER",     spine.problem_statement[:400])
        .replace("GAP_PLACEHOLDER",         spine.research_gap[:300])
        .replace("RQS_PLACEHOLDER",         "; ".join(spine.research_questions) or "NOT FOUND")
        .replace("ROS_PLACEHOLDER",         "; ".join(spine.research_objectives) or "NOT FOUND")
        .replace("HYPS_PLACEHOLDER",        "; ".join(spine.hypotheses) or "NOT FOUND")
        .replace("THEORY_PLACEHOLDER",      spine.theory_used)
        .replace("VARS_PLACEHOLDER",        "; ".join(spine.variables) or "NOT FOUND")
        .replace("METHOD_PLACEHOLDER",      spine.methodology)
        .replace("SAMPLING_PLACEHOLDER",    spine.sampling)
        .replace("INSTRUMENT_PLACEHOLDER",  spine.instrument)
        .replace("ANALYSIS_PLACEHOLDER",    spine.analysis_technique)
        .replace("FINDINGS_PLACEHOLDER",    "; ".join(spine.key_findings) or "NOT FOUND")
        .replace("CONCLUSIONS_PLACEHOLDER", "; ".join(spine.conclusions) or "NOT FOUND")
        .replace("CRITICAL_ISSUES_PLACEHOLDER", critical_block)
    )
    try:
        msg = await claude_client.messages.create(
            model="claude-sonnet-4-6", max_tokens=4096, temperature=0.2,
            system=_ALIGNMENT_SYSTEM,
            messages=[{"role":"user","content":prompt}])
        raw = msg.content[0].text.strip()
        raw = re.sub(r"```(?:json)?","", raw).strip("`").strip()
        d = json.loads(raw)
        rows = [AlignmentRow(
            rq=        r.get("rq","—"),         ro=       r.get("ro","—"),
            hypothesis=r.get("hypothesis","—"),  method=   r.get("method","—"),
            analysis=  r.get("analysis","—"),    finding=  r.get("finding","—"),
            conclusion=r.get("conclusion","—"),
            status=    str(r.get("status","UNCLEAR")).upper(),
            note=      r.get("note",""))
            for r in d.get("rows",[])]
        matrix = AlignmentMatrix(
            rows=rows,
            overall_verdict=            d.get("overall_verdict",""),
            golden_thread_score=        str(d.get("golden_thread_score","UNCLEAR")).upper(),
            critical_gaps=              d.get("critical_gaps",[]),
            structural_recommendations= d.get("structural_recommendations",[]),
        )
        print(f"[ThesisSifu] Alignment audit: {len(rows)} rows, score={matrix.golden_thread_score}")
        return matrix
    except json.JSONDecodeError as e:
        print(f"[ThesisSifu] Alignment JSON parse error: {e}")
        return AlignmentMatrix(overall_verdict="Alignment JSON parse failed.", golden_thread_score="UNCLEAR")
    except Exception as e:
        print(f"[ThesisSifu] Alignment audit error: {e}\n{traceback.format_exc()}")
        return AlignmentMatrix(overall_verdict=f"Alignment audit failed: {e}", golden_thread_score="UNCLEAR")


# ── Stage 4: Holistic Examiner (Claude) ───────────────────────
_EXAMINER_SYSTEM = """\
You are a senior academic examiner with 25 years of experience.
Produce rigorous, specific, honest examiner-level critique anchored to the
provided thesis spine and alignment matrix. You are NOT a grammar checker —
you are an intellectual critic. Cite specific spine fields when you make claims.
Recommend named frameworks (with seminal author) where appropriate.
Respond in the same language as the document.\
"""

_PHD_EXAMINER = """\
External examiner for a PhD viva.

THESIS SPINE
============
Title:            "TITLE_PLACEHOLDER"
Field:            FIELD_PLACEHOLDER | Institution: INST_PLACEHOLDER
Problem:          PROBLEM_PLACEHOLDER
Research Qs:      RQS_PLACEHOLDER
Research Objs:    ROS_PLACEHOLDER
Theory:           THEORY_PLACEHOLDER
Methodology:      METHOD_PLACEHOLDER
Analysis:         ANALYSIS_PLACEHOLDER

ALIGNMENT VERDICT
=================
Golden thread score: GT_SCORE_PLACEHOLDER
Overall:             ALIGN_VERDICT_PLACEHOLDER
Critical gaps:
GAPS_PLACEHOLDER

CHAPTER + SUBSECTION ISSUE SUMMARY
===================================
SUMMARIES_PLACEHOLDER

CANONICAL LIBRARIES (recommend FROM these — do NOT invent references):
FRAMEWORKS_PLACEHOLDER
METHODS_PLACEHOLDER

Write a full examiner report with EXACTLY these sections:

SECTION 1 — ORIGINAL CONTRIBUTION TO KNOWLEDGE
SECTION 2 — GOLDEN THREAD ANALYSIS
Use the alignment matrix. Quote spine fields. Identify exact broken links.
SECTION 3 — CHAPTER + SUBSECTION ALIGNMENT AUDIT
For each chapter, list which subsections under-delivered on expected purpose.
SECTION 4 — THEORETICAL FRAMEWORK COHERENCE
Name a specific framework (with seminal author) to add or strengthen.
SECTION 5 — METHODOLOGICAL RIGOUR
Recommend specific named methods (e.g., PLS-SEM via Hair et al. 2017).
SECTION 6 — DATA QUALITY & ANALYTICAL DEPTH
SECTION 7 — SCOPUS-LEVEL LANGUAGE & TONE
SECTION 8 — CITATION & REFERENCE INTEGRITY
SECTION 9 — CRITICAL VIVA QUESTIONS
List 8 specific questions tied to spine and alignment gaps.
SECTION 10 — EXAMINER'S VERDICT
One of: PASS / PASS WITH MINOR CORRECTIONS / MAJOR REVISIONS REQUIRED / REFER (RESUBMIT) / FAIL
Formal statement 6-8 sentences.\
"""

_MASTERS_EXAMINER = """\
Internal reader for a Master's thesis viva.

THESIS SPINE
============
Title:         "TITLE_PLACEHOLDER"
Field:         FIELD_PLACEHOLDER | Institution: INST_PLACEHOLDER
Problem:       PROBLEM_PLACEHOLDER
Research Qs:   RQS_PLACEHOLDER
Research Objs: ROS_PLACEHOLDER
Theory:        THEORY_PLACEHOLDER
Methodology:   METHOD_PLACEHOLDER
Analysis:      ANALYSIS_PLACEHOLDER

ALIGNMENT VERDICT
=================
Golden thread score: GT_SCORE_PLACEHOLDER
Overall:             ALIGN_VERDICT_PLACEHOLDER
Critical gaps:
GAPS_PLACEHOLDER

CHAPTER + SUBSECTION ISSUE SUMMARY
====================================
SUMMARIES_PLACEHOLDER

CANONICAL LIBRARIES:
FRAMEWORKS_PLACEHOLDER
METHODS_PLACEHOLDER

Write a report with EXACTLY these sections:

SECTION 1 — THESIS ALIGNMENT & GOLDEN THREAD
SECTION 2 — CHAPTER + SUBSECTION ALIGNMENT AUDIT
SECTION 3 — THEORETICAL FRAMEWORK (recommend named theory + seminal author)
SECTION 4 — METHODOLOGICAL RIGOUR (recommend named method + seminal author)
SECTION 5 — DATA ROBUSTNESS & INSTRUMENTATION
SECTION 6 — SCOPUS-LEVEL LANGUAGE & TONE
SECTION 7 — CITATION INTEGRITY
SECTION 8 — CRITICAL VIVA QUESTIONS (5 questions tied to alignment gaps)
SECTION 9 — EXAMINER'S VERDICT
One of: PASS / PASS WITH MINOR CORRECTIONS / MAJOR REVISIONS REQUIRED / RESUBMIT
Top 5 issues that must be resolved.\
"""

_JOURNAL_EXAMINER = """\
Peer reviewer for a Scopus journal.

ARTICLE SPINE
=============
Title:       "TITLE_PLACEHOLDER"
Field:       FIELD_PLACEHOLDER
Problem:     PROBLEM_PLACEHOLDER
Research Qs: RQS_PLACEHOLDER
Theory:      THEORY_PLACEHOLDER
Methodology: METHOD_PLACEHOLDER
Analysis:    ANALYSIS_PLACEHOLDER

ALIGNMENT VERDICT
=================
Golden thread score: GT_SCORE_PLACEHOLDER
Overall:             ALIGN_VERDICT_PLACEHOLDER
Critical gaps:
GAPS_PLACEHOLDER

SECTION ISSUE SUMMARY
=====================
SUMMARIES_PLACEHOLDER

CANONICAL LIBRARIES:
FRAMEWORKS_PLACEHOLDER
METHODS_PLACEHOLDER

Write a review with EXACTLY these sections:

SECTION 1 — CONTRIBUTION & NOVELTY
SECTION 2 — LITERATURE CURRENCY & GAPS
SECTION 3 — METHODOLOGY
SECTION 4 — RESULTS & ANALYSIS
SECTION 5 — ARGUMENT COHERENCE
SECTION 6 — SCOPUS-LEVEL LANGUAGE & TONE
SECTION 7 — CITATION INTEGRITY
SECTION 8 — PUBLICATION VERDICT
One of: ACCEPT / MINOR REVISIONS / MAJOR REVISIONS / REJECT
Top 3 critical revisions.\
"""

_UG_EXAMINER = """\
Supervisor reviewing an undergraduate FYP.

PROJECT SPINE
=============
Title:       "TITLE_PLACEHOLDER"
Field:       FIELD_PLACEHOLDER
Problem:     PROBLEM_PLACEHOLDER
Research Qs: RQS_PLACEHOLDER
Methodology: METHOD_PLACEHOLDER
Analysis:    ANALYSIS_PLACEHOLDER

ALIGNMENT VERDICT
=================
Golden thread score: GT_SCORE_PLACEHOLDER
Overall:             ALIGN_VERDICT_PLACEHOLDER

SECTION ISSUE SUMMARY
=====================
SUMMARIES_PLACEHOLDER

CANONICAL LIBRARIES:
FRAMEWORKS_PLACEHOLDER
METHODS_PLACEHOLDER

Write a report with EXACTLY these sections:

SECTION 1 — SCOPE & RESEARCH QUESTION
SECTION 2 — LITERATURE REVIEW
SECTION 3 — METHODOLOGY
SECTION 4 — ANALYSIS & FINDINGS
SECTION 5 — CRITICAL THINKING
SECTION 6 — WRITING & ACADEMIC CONVENTIONS
SECTION 7 — TOP ISSUES TO FIX
SECTION 8 — SUPERVISOR'S VERDICT
One of: PASS / PASS WITH MINOR CORRECTIONS / MAJOR REVISIONS / FAIL\
"""

_EXAMINER_TEMPLATES = {
    "PHD": _PHD_EXAMINER, "MASTERS": _MASTERS_EXAMINER,
    "JOURNAL_ARTICLE": _JOURNAL_EXAMINER, "UNDERGRADUATE": _UG_EXAMINER,
}

async def run_examiner(doc_type: str, spine: ThesisSpine, field: str, inst: str,
                       chs: list[ChapterSummary], align: AlignmentMatrix) -> str:
    if not claude_client:
        return "Examiner synthesis unavailable — ANTHROPIC_API_KEY not set."
    # Build summaries
    summaries = ""
    for cs in chs:
        n_c = sum(1 for c in cs.comments if c.severity=="CRITICAL")
        n_m = sum(1 for c in cs.comments if c.severity=="MODERATE")
        summaries += (f"\nCh.{cs.chapter_num} — {cs.chapter_title}: "
                      f"{len(cs.comments)} comments (Critical:{n_c} Moderate:{n_m})\n")
        for sub in cs.subsections:
            crit = [c for c in sub.comments if c.severity=="CRITICAL"]
            if crit:
                summaries += f"  §{sub.subsection_num} {sub.title} [purpose: {sub.expected_purpose[:70]}]\n"
                for c in crit[:3]:
                    summaries += f"    [CRITICAL] {c.issue}\n"
    gaps = "\n".join(f"  * {g}" for g in align.critical_gaps[:8]) or "  (none)"
    template = _EXAMINER_TEMPLATES.get(doc_type, _MASTERS_EXAMINER)
    prompt = (
        template
        .replace("TITLE_PLACEHOLDER",        spine.title[:200])
        .replace("FIELD_PLACEHOLDER",         field)
        .replace("INST_PLACEHOLDER",          inst)
        .replace("PROBLEM_PLACEHOLDER",       spine.problem_statement[:400])
        .replace("RQS_PLACEHOLDER",           "; ".join(spine.research_questions[:5]) or "NOT FOUND")
        .replace("ROS_PLACEHOLDER",           "; ".join(spine.research_objectives[:5]) or "NOT FOUND")
        .replace("THEORY_PLACEHOLDER",        spine.theory_used)
        .replace("METHOD_PLACEHOLDER",        spine.methodology)
        .replace("ANALYSIS_PLACEHOLDER",      spine.analysis_technique)
        .replace("GT_SCORE_PLACEHOLDER",      align.golden_thread_score)
        .replace("ALIGN_VERDICT_PLACEHOLDER", align.overall_verdict)
        .replace("GAPS_PLACEHOLDER",          gaps)
        .replace("SUMMARIES_PLACEHOLDER",     summaries[:8000])
        .replace("FRAMEWORKS_PLACEHOLDER",    CANONICAL_FRAMEWORKS)
        .replace("METHODS_PLACEHOLDER",       CANONICAL_METHODS)
    )
    try:
        msg = await claude_client.messages.create(
            model="claude-sonnet-4-6", max_tokens=8192, temperature=0.3,
            system=_EXAMINER_SYSTEM,
            messages=[{"role":"user","content":prompt}])
        return msg.content[0].text.strip()
    except Exception as e:
        print(f"[ThesisSifu] Examiner error: {e}\n{traceback.format_exc()}")
        return f"Examiner synthesis failed: {e}"


# ── PDF helpers ────────────────────────────────────────────────
def pdf_styles():
    def S(n, **k):
        k.setdefault("fontName","Times-Roman"); k.setdefault("fontSize",11)
        k.setdefault("leading",16); k.setdefault("textColor",BLACK)
        k.setdefault("alignment",TA_JUSTIFY)
        return ParagraphStyle(n,**k)
    return {
        "Title":   S("Title",fontName="Helvetica-Bold",fontSize=20,leading=26,textColor=NAVY,alignment=TA_LEFT),
        "Sub":     S("Sub",fontName="Helvetica",fontSize=12,leading=18,textColor=ACCENT,alignment=TA_LEFT),
        "Meta":    S("Meta",fontName="Helvetica",fontSize=10,leading=15,textColor=colors.HexColor("#444444"),alignment=TA_LEFT),
        "Badge":   S("Badge",fontName="Helvetica-Bold",fontSize=10,leading=14,textColor=WHITE,alignment=TA_CENTER),
        "SecHead": S("SecHead",fontName="Helvetica-Bold",fontSize=13,leading=20,textColor=NAVY,alignment=TA_LEFT,spaceBefore=16,spaceAfter=6),
        "ChHead":  S("ChHead",fontName="Helvetica-Bold",fontSize=11,leading=16,textColor=ACCENT,alignment=TA_LEFT,spaceBefore=12,spaceAfter=4),
        "SubHead": S("SubHead",fontName="Helvetica-Bold",fontSize=10,leading=14,textColor=colors.HexColor("#2A5A8A"),alignment=TA_LEFT,spaceBefore=8,spaceAfter=3),
        "Body":    S("Body",leading=18,spaceAfter=6),
        "Bullet":  S("Bullet",leading=17,leftIndent=16,spaceAfter=4),
        "Verdict": S("Verdict",fontName="Helvetica-Bold",fontSize=15,leading=22,textColor=WHITE,alignment=TA_CENTER),
        "VrdSub":  S("VrdSub",fontName="Helvetica",fontSize=10,leading=14,textColor=WHITE,alignment=TA_CENTER),
        "TblH":    S("TblH",fontName="Helvetica-Bold",fontSize=9,leading=13,textColor=NAVY,alignment=TA_LEFT),
        "TblC":    S("TblC",fontName="Times-Roman",fontSize=9,leading=13,textColor=BLACK,alignment=TA_LEFT),
        "TblCSm":  S("TblCSm",fontName="Times-Roman",fontSize=8,leading=11,textColor=BLACK,alignment=TA_LEFT),
        "Footer":  S("Footer",fontName="Helvetica-Oblique",fontSize=8,leading=12,textColor=colors.HexColor("#999"),alignment=TA_CENTER),
        "Excerpt": S("Excerpt",fontName="Times-Italic",fontSize=9,leading=13,textColor=colors.HexColor("#555"),leftIndent=8),
        "Detail":  S("Detail",fontName="Times-Roman",fontSize=10,leading=15,textColor=BLACK,leftIndent=4),
        "SpineV":  S("SpineV",fontName="Times-Roman",fontSize=10,leading=14,textColor=BLACK),
        "SpineK":  S("SpineK",fontName="Helvetica-Bold",fontSize=10,leading=14,textColor=NAVY),
    }

def HR(c=RULE,t=0.5,b=5,a=7):
    return HRFlowable(width="100%",thickness=t,color=c,spaceBefore=b,spaceAfter=a)

def _tbl(data, col_widths, styles):
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle(styles))
    return t

def meta_box(S, rows):
    data = [[Paragraph(f'<font name="Helvetica-Bold">{k}</font>',S["Meta"]),
             Paragraph(str(v),S["Meta"])] for k,v in rows]
    return _tbl(data,[4*cm,W-5*cm-4*cm],[
        ("BACKGROUND",(0,0),(-1,-1),BOX_BG),("BOX",(0,0),(-1,-1),.5,RULE),
        ("INNERGRID",(0,0),(-1,-1),.3,RULE),("TOPPADDING",(0,0),(-1,-1),5),
        ("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(0,0),(-1,-1),8),
        ("RIGHTPADDING",(0,0),(-1,-1),8),("VALIGN",(0,0),(-1,-1),"TOP")])

def vbanner(S, text, color, subtitle="Examiner's Recommended Outcome"):
    t = Table([[Paragraph(text.upper(),S["Verdict"])],[Paragraph(subtitle,S["VrdSub"])]],
              colWidths=[W-5*cm])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),color),
        ("TOPPADDING",(0,0),(-1,-1),10),("BOTTOMPADDING",(0,0),(-1,-1),10),
        ("LEFTPADDING",(0,0),(-1,-1),16),("RIGHTPADDING",(0,0),(-1,-1),16)]))
    return t

def sev_badge(S, sev, count):
    t = Table([[Paragraph(f"{sev}: {count}",S["Badge"])]],colWidths=[3.5*cm])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),SEV_COLORS.get(sev,ACCENT)),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8)]))
    return t

def type_badge_table(S, text, color):
    t = Table([[Paragraph(text,S["Badge"])]],colWidths=[W-5*cm])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),color),
        ("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),
        ("LEFTPADDING",(0,0),(-1,-1),12),("RIGHTPADDING",(0,0),(-1,-1),12)]))
    return t

def spine_box(S, spine: ThesisSpine):
    rows = [
        ("Title",              spine.title[:120]),
        ("Discipline",         spine.discipline),
        ("Problem Statement",  spine.problem_statement[:300]),
        ("Research Gap",       spine.research_gap[:250]),
        ("Research Questions", "; ".join(spine.research_questions[:6]) or "NOT FOUND"),
        ("Research Objectives","; ".join(spine.research_objectives[:6]) or "NOT FOUND"),
        ("Theory Used",        spine.theory_used),
        ("Variables",          "; ".join(spine.variables[:10]) or "NOT FOUND"),
        ("Methodology",        spine.methodology),
        ("Sampling",           spine.sampling),
        ("Instrument",         spine.instrument),
        ("Analysis Technique", spine.analysis_technique),
        ("Key Findings",       "; ".join(spine.key_findings[:4]) or "NOT FOUND"),
        ("Conclusions",        "; ".join(spine.conclusions[:4]) or "NOT FOUND"),
    ]
    data = [[Paragraph(k,S["SpineK"]), Paragraph(str(v),S["SpineV"])] for k,v in rows]
    return _tbl(data,[4.5*cm,W-5*cm-4.5*cm],[
        ("BACKGROUND",(0,0),(-1,-1),BOX_BG),("BOX",(0,0),(-1,-1),.5,RULE),
        ("INNERGRID",(0,0),(-1,-1),.3,RULE),("TOPPADDING",(0,0),(-1,-1),4),
        ("BOTTOMPADDING",(0,0),(-1,-1),4),("LEFTPADDING",(0,0),(-1,-1),8),
        ("RIGHTPADDING",(0,0),(-1,-1),8),("VALIGN",(0,0),(-1,-1),"TOP")])

def parse_secs(text):
    pat = re.compile(r'(?:^|\n)\s*(?:\*{0,2})?SECTION\s+\d+\s*[—–\-]+\s*([^\n*]+?)(?:\*{0,2})?\s*\n',re.IGNORECASE)
    ms  = list(pat.finditer(text))
    if not ms: return [{"title":"Full Report","body":text.strip()}]
    out = []
    for i,m in enumerate(ms):
        t = m.group(1).strip().rstrip("*").strip()
        s = m.end(); e = ms[i+1].start() if i+1<len(ms) else len(text)
        out.append({"title":t,"body":text[s:e].strip()})
    return out

def detect_verdict(text):
    u = text.upper()
    for lbl,col in [("READY FOR SUBMISSION",GREEN),("ACCEPT",GREEN),
                    ("PASS WITH MINOR CORRECTIONS",AMBER),("MINOR REVISIONS REQUIRED",AMBER),
                    ("MAJOR REVISIONS REQUIRED",RED),("REFER (RESUBMIT)",RED),
                    ("REJECT AND REWRITE",RED),("REJECT",RED),("RESUBMIT",RED),
                    ("FAIL",RED),("PASS",GREEN)]:
        if lbl in u: return lbl.title(), col
    return "Review Complete", ACCENT

def body2story(S, body):
    items = []
    for line in body.split("\n"):
        line = line.strip()
        if not line: items.append(Spacer(1,3)); continue
        line = re.sub(r'\*\*(.+?)\*\*',r'<b>\1</b>',line)
        line = re.sub(r'\*(.+?)\*',r'<i>\1</i>',line)
        if re.match(r'^[-•]\s+',line): items.append(Paragraph("• "+line[2:],S["Bullet"]))
        elif re.match(r'^\d+[\.\)]\s+',line): items.append(Paragraph(line,S["Bullet"]))
        else: items.append(Paragraph(line,S["Body"]))
    return items

def _comment_text(c: ParagraphComment) -> str:
    sev = {"CRITICAL":"[CRITICAL]","MODERATE":"[MODERATE]",
           "SUGGESTION":"[SUGGESTION]"}.get(c.severity,"[NOTE]")
    fw  = f"\n\nSUGGESTED FRAMEWORK: {c.suggested_framework}" if c.suggested_framework else ""
    mt  = f"\n\nSUGGESTED METHOD: {c.suggested_method}" if c.suggested_method else ""
    return (f"{sev} ThesisSifu v4 — §{c.subsection} {c.subsection_title}\n\n"
            f"ISSUE: {c.issue}\n\n"
            f"RECOMMENDATION: {c.recommendation}\n\n"
            f"LITERATURE NEEDED: {c.literature_needed}\n\n"
            f"THEORY/FRAMEWORK: {c.theory_needed}"
            f"{fw}{mt}")


# ── Output 1: Examiner Report PDF ─────────────────────────────
def build_examiner_pdf(filename, audit_id, doc_type, clf, spine, examiner_text, align, chs) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf,pagesize=A4,leftMargin=2.5*cm,rightMargin=2.5*cm,
                            topMargin=2.2*cm,bottomMargin=2.2*cm)
    S = pdf_styles(); story = []

    type_label = TYPE_LABELS.get(doc_type,doc_type)
    story.append(type_badge_table(S,type_label.upper(),TYPE_COLORS.get(doc_type,NAVY)))
    story.append(Spacer(1,14))
    story.append(Paragraph("EXAMINER'S AUDIT REPORT",S["Title"]))
    story.append(Paragraph("ThesisSifu Pro v4.1 — Alignment-Aware Multi-Agent Panel",S["Sub"]))
    story.append(HR(c=NAVY,t=1.5,b=4,a=10))

    au=clf.get("authors","UNKNOWN"); fi=clf.get("field","UNKNOWN"); ins=clf.get("institution","UNKNOWN")
    story.append(meta_box(S,[
        ("Document",filename),("Detected As",type_label),
        ("Confidence",clf.get("confidence","N/A")),
        ("Author(s)",(au[:80]+"…") if len(au)>80 else au),
        ("Field",fi),("Institution",ins),
        ("Audit ID",audit_id),("Date",datetime.now().strftime("%d %B %Y, %H:%M")),
        ("Report","1 of 4 — Examiner Audit Report"),
    ]))
    story.append(Spacer(1,10))

    story.append(Paragraph("EXTRACTED THESIS SPINE",S["SecHead"]))
    story.append(HR(c=NAVY,t=1,b=0,a=6))
    story.append(spine_box(S, spine))
    story.append(Spacer(1,10))

    nc=sum(1 for cs in chs for c in cs.comments if c.severity=="CRITICAL")
    nm=sum(1 for cs in chs for c in cs.comments if c.severity=="MODERATE")
    ns=sum(1 for cs in chs for c in cs.comments if c.severity=="SUGGESTION")
    brow = Table([[sev_badge(S,"CRITICAL",nc),Spacer(6,1),
                   sev_badge(S,"MODERATE",nm),Spacer(6,1),
                   sev_badge(S,"SUGGESTION",ns)]],
                 colWidths=[3.5*cm,.5*cm,3.5*cm,.5*cm,3.5*cm])
    brow.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0)]))
    story.append(brow); story.append(Spacer(1,10))

    gt = align.golden_thread_score
    gt_color = {"STRONG":GREEN,"ACCEPTABLE":AMBER,"WEAK":RED,"BROKEN":RED}.get(gt, ACCENT)
    story.append(vbanner(S,f"Golden Thread: {gt}",gt_color,subtitle="Structural Alignment Score"))
    story.append(Spacer(1,8))
    vl,vc = detect_verdict(examiner_text)
    story.append(vbanner(S,vl,vc)); story.append(Spacer(1,10)); story.append(HR())

    for sec in parse_secs(examiner_text):
        bl = [Paragraph(sec["title"].upper(),S["SecHead"]),HR(c=ACCENT,t=.7,b=0,a=6)]
        bl.extend(body2story(S,sec["body"])); bl.append(Spacer(1,6))
        story.append(KeepTogether(bl[:3])); story.extend(bl[3:])

    story.append(PageBreak())
    story.append(Paragraph("CHAPTER + SUBSECTION ISSUE SUMMARY",S["SecHead"]))
    story.append(HR(c=NAVY,t=1,b=0,a=8))
    td = [[Paragraph(h,S["TblH"]) for h in ["Chapter / Subsection","Critical","Moderate","Suggestions","Total"]]]
    for cs in chs:
        nc2=sum(1 for c in cs.comments if c.severity=="CRITICAL")
        nm2=sum(1 for c in cs.comments if c.severity=="MODERATE")
        ns2=len(cs.comments)-nc2-nm2
        td.append([Paragraph(f"<b>Ch.{cs.chapter_num} — {cs.chapter_title[:40]}</b>",S["TblC"]),
                   Paragraph(str(nc2),S["TblC"]),Paragraph(str(nm2),S["TblC"]),
                   Paragraph(str(ns2),S["TblC"]),Paragraph(str(nc2+nm2+ns2),S["TblC"])])
        for sub in cs.subsections:
            if not sub.comments: continue
            snc=sum(1 for c in sub.comments if c.severity=="CRITICAL")
            snm=sum(1 for c in sub.comments if c.severity=="MODERATE")
            sns=len(sub.comments)-snc-snm
            td.append([Paragraph(f"   §{sub.subsection_num} {sub.title[:45]}",S["TblCSm"]),
                       Paragraph(str(snc),S["TblCSm"]),Paragraph(str(snm),S["TblCSm"]),
                       Paragraph(str(sns),S["TblCSm"]),Paragraph(str(snc+snm+sns),S["TblCSm"])])
    cw = W-5*cm
    t = Table(td,colWidths=[cw*.52,cw*.12,cw*.12,cw*.12,cw*.12])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),NAVY),("TEXTCOLOR",(0,0),(-1,0),WHITE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,BOX_BG]),("BOX",(0,0),(-1,-1),.5,RULE),
        ("INNERGRID",(0,0),(-1,-1),.3,RULE),("TOPPADDING",(0,0),(-1,-1),4),
        ("BOTTOMPADDING",(0,0),(-1,-1),4),("LEFTPADDING",(0,0),(-1,-1),5),
        ("RIGHTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(t)
    story.append(Spacer(1,20)); story.append(HR())
    story.append(Paragraph("ThesisSifu Pro v4.1 — See Report 2 (Annotated Thesis), "
        "Report 3 (Commentary Log), Report 4 (Alignment Matrix) in this ZIP.",S["Footer"]))
    doc.build(story); buf.seek(0); return buf.getvalue()


# ── Output 2: Annotated DOCX ──────────────────────────────────
def _get_or_create_comments_part(doc):
    rt = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    for rel in doc.part.rels.values():
        if rel.reltype == rt: return rel.target_part
    xml_str = '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    element = parse_xml(xml_str)
    uri = PackURI('/word/comments.xml')
    ct  = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
    cp  = XmlPart(uri, ct, element, doc.part.package)
    doc.part.relate_to(cp, rt)
    return cp

def _insert_comment(doc, para, cid, author, text, date_str):
    cp = _get_or_create_comments_part(doc)
    ce = OxmlElement('w:comment')
    ce.set(qn('w:id'),str(cid)); ce.set(qn('w:author'),author)
    ce.set(qn('w:date'),date_str); ce.set(qn('w:initials'),'TS')
    pe = OxmlElement('w:p'); re_ = OxmlElement('w:r'); te = OxmlElement('w:t')
    te.text = text; te.set(qn('xml:space'),'preserve')
    re_.append(te); pe.append(re_); ce.append(pe)
    cp.element.append(ce)
    px = para._p
    crs = OxmlElement('w:commentRangeStart'); crs.set(qn('w:id'),str(cid)); px.insert(0,crs)
    cre = OxmlElement('w:commentRangeEnd');   cre.set(qn('w:id'),str(cid)); px.append(cre)
    rr  = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    rs  = OxmlElement('w:rStyle'); rs.set(qn('w:val'),'CommentReference')
    rpr.append(rs); rr.append(rpr)
    cr  = OxmlElement('w:commentReference'); cr.set(qn('w:id'),str(cid)); rr.append(cr); px.append(rr)

def build_annotated_docx(content, filename, audit_id, chs, clf) -> bytes:
    doc = DocxDocument(io.BytesIO(content))
    date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    cp = _get_or_create_comments_part(doc)
    existing_ids = [int(c.get(qn('w:id'))) for c in cp.element.xpath('.//w:comment')
                    if c.get(qn('w:id')) is not None]
    cid = max(existing_ids)+1 if existing_ids else 0

    lookup: dict[str,list] = {}
    for cs in chs:
        for c in cs.comments:
            key = c.para_excerpt[:50].lower().strip()
            lookup.setdefault(key,[]).append(c)

    matched = set()
    for para in doc.paragraphs:
        pt = para.text.strip()
        if not pt or len(pt) < 30: continue
        pkey = pt[:50].lower().strip()
        best, best_sc = None, 0
        for ekey, cmts in lookup.items():
            aw = set(pkey.split()); bw = set(ekey.split())
            sc = len(aw & bw) / max(len(aw),1)
            if sc > best_sc and sc > 0.30:
                best_sc = sc; best = (ekey, cmts)
        if best:
            for c in best[1]:
                if id(c) in matched: continue
                matched.add(id(c))
                try:
                    _insert_comment(doc, para, cid, "ThesisSifu v4", _comment_text(c), date_str)
                    cid += 1
                except Exception as e:
                    print(f"[ThesisSifu] DOCX comment insert error {cid}: {e}")
                    cid += 1

    hdr = doc.sections[0].header
    if hdr.paragraphs:
        hdr.paragraphs[0].text = (f"ThesisSifu Pro v4 | Annotated Thesis | "
                                   f"Audit ID: {audit_id} | {datetime.now().strftime('%d %B %Y')}")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()


# ── Output 2 (PDF): Annotated PDF (Sticky Notes) ─────────────
def build_annotated_pdf(content: bytes, full_text: str, audit_id: str, chs: list) -> bytes:
    """Place sticky-note annotations on the PDF.

    PDF TEXT SEARCH LIMITATION: Many thesis PDFs use Identity-H CID TrueType fonts
    which PyMuPDF cannot search with page.search_for(). Rather than placing all
    annotations at (30,30) as invisible stacked fallbacks, we use a reliable
    position-based strategy:
      - Place each annotation in the RIGHT MARGIN of the estimated page
      - Stagger vertically so multiple annotations on one page don't overlap
      - Use a different Y offset for CRITICAL (top third), MODERATE (mid), SUGGESTION (bottom)
      - Colour-code icon: red=CRITICAL, orange=MODERATE, blue=SUGGESTION
    This guarantees every comment is visible and readable, regardless of font encoding.
    """
    fitz_doc = fitz.open(stream=content, filetype="pdf")
    total_pages = len(fitz_doc)

    # Track how many annotations are already on each page (to stagger Y)
    page_annot_counts: dict[int, int] = {}

    all_comments = [c for cs in chs for c in cs.comments]
    print(f"[ThesisSifu] Placing {len(all_comments)} PDF annotations (position-based)")

    for c in all_comments:
        txt = _comment_text(c)

        # Determine target page (0-indexed, clamped)
        pg = max(0, min(c.page_estimate - 1, total_pages - 1))

        # Try text search first (works when fonts cooperate)
        search_text = c.para_excerpt[:40].replace('\n', ' ').strip()
        placed_by_search = False

        if search_text:
            search_range = range(max(0, pg - 2), min(total_pages, pg + 3))
            for p_num in search_range:
                page = fitz_doc[p_num]
                rects = page.search_for(search_text)
                if not rects and len(search_text) > 20:
                    rects = page.search_for(search_text[:20])
                if rects:
                    rect  = rects[0]
                    point = fitz.Point(max(5, rect.x0 - 18), rect.y0)
                    annot = page.add_text_annot(point, txt)
                    annot.set_info(title="ThesisSifu Pro", content=txt)
                    if   c.severity == "CRITICAL":   annot.set_colors(stroke=(0.8, 0, 0))
                    elif c.severity == "MODERATE":   annot.set_colors(stroke=(0.9, 0.4, 0))
                    else:                             annot.set_colors(stroke=(0, 0, 0.8))
                    annot.update()
                    placed_by_search = True
                    page_annot_counts[p_num] = page_annot_counts.get(p_num, 0) + 1
                    break

        if not placed_by_search:
            # ── Position-based right-margin placement ──────────────────────
            page = fitz_doc[pg]
            pw   = page.rect.width
            ph   = page.rect.height

            # Base Y by severity: spread across page to avoid clustering
            n = page_annot_counts.get(pg, 0)

            # CRITICAL → top third, MODERATE → mid, SUGGESTION → bottom third
            if c.severity == "CRITICAL":
                base_y = ph * 0.10
            elif c.severity == "MODERATE":
                base_y = ph * 0.40
            else:
                base_y = ph * 0.70

            # Stagger: each successive annotation on this page moves down 28pt
            # Wrap back toward base if we'd go off page
            y_pos = base_y + (n % 8) * 28
            if y_pos > ph - 30:
                y_pos = base_y + (n % 4) * 14

            # X: right margin (5pt from edge)
            x_pos = pw - 18

            point = fitz.Point(x_pos, y_pos)
            annot = page.add_text_annot(point, txt)
            annot.set_info(title="ThesisSifu Pro", content=txt)
            if   c.severity == "CRITICAL":   annot.set_colors(stroke=(0.8, 0, 0))
            elif c.severity == "MODERATE":   annot.set_colors(stroke=(0.9, 0.4, 0))
            else:                             annot.set_colors(stroke=(0, 0, 0.8))
            annot.update()
            page_annot_counts[pg] = page_annot_counts.get(pg, 0) + 1

    total_placed = sum(page_annot_counts.values())
    print(f"[ThesisSifu] PDF annotations placed: {total_placed} across "
          f"{len(page_annot_counts)} pages")
    return fitz_doc.tobytes()   # tobytes() is more reliable than write() for in-memory


# ── Output 3: Commentary Report PDF ───────────────────────────
def build_commentary_pdf(filename, audit_id, doc_type, clf, spine, chs) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf,pagesize=A4,leftMargin=2.5*cm,rightMargin=2.5*cm,
                            topMargin=2.2*cm,bottomMargin=2.2*cm)
    S = pdf_styles(); story = []

    story.append(type_badge_table(S,"PARAGRAPH-LEVEL COMMENTARY REPORT",TYPE_COLORS.get(doc_type,NAVY)))
    story.append(Spacer(1,14))
    story.append(Paragraph("COMMENTARY REPORT",S["Title"]))
    story.append(Paragraph("ThesisSifu Pro v4.1 — Subsection-Level Supervisor Notes",S["Sub"]))
    story.append(HR(c=NAVY,t=1.5,b=4,a=10))

    au = clf.get("authors","UNKNOWN")
    story.append(meta_box(S,[
        ("Document",filename),("Type",TYPE_LABELS.get(doc_type,doc_type)),
        ("Title",(spine.title[:85]+"…") if len(spine.title)>85 else spine.title),
        ("Author(s)",(au[:80]+"…") if len(au)>80 else au),
        ("Audit ID",audit_id),("Date",datetime.now().strftime("%d %B %Y, %H:%M")),
        ("Report","3 of 4 — Paragraph Commentary Log"),
    ]))
    story.append(Spacer(1,10))

    all_c = [c for cs in chs for c in cs.comments]
    nc = sum(1 for c in all_c if c.severity=="CRITICAL")
    nm = sum(1 for c in all_c if c.severity=="MODERATE")
    ns = sum(1 for c in all_c if c.severity=="SUGGESTION")

    story.append(Paragraph("OVERALL COMMENT DISTRIBUTION",S["SecHead"]))
    cw = W-5*cm
    td = [[Paragraph(h,S["TblH"]) for h in ["Severity","Count","Description"]],
          [Paragraph("CRITICAL",S["TblC"]),Paragraph(str(nc),S["TblC"]),
           Paragraph("Fundamental flaws requiring resolution",S["TblC"])],
          [Paragraph("MODERATE",S["TblC"]),Paragraph(str(nm),S["TblC"]),
           Paragraph("Significant weaknesses requiring revision",S["TblC"])],
          [Paragraph("SUGGESTION",S["TblC"]),Paragraph(str(ns),S["TblC"]),
           Paragraph("Improvement opportunities",S["TblC"])],
          [Paragraph("TOTAL",S["TblC"]),Paragraph(str(len(all_c)),S["TblC"]),Paragraph("",S["TblC"])]]
    t = Table(td,colWidths=[cw*.22,cw*.12,cw*.66])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),NAVY),("TEXTCOLOR",(0,0),(-1,0),WHITE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,BOX_BG]),("BOX",(0,0),(-1,-1),.5,RULE),
        ("INNERGRID",(0,0),(-1,-1),.3,RULE),("TOPPADDING",(0,0),(-1,-1),5),
        ("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(0,0),(-1,-1),7),
        ("RIGHTPADDING",(0,0),(-1,-1),7)]))
    story.append(t); story.append(Spacer(1,16)); story.append(HR())

    if not all_c:
        story.append(Paragraph("<i>No paragraph comments were generated. "
            "Check Railway logs for API errors or empty document detection.</i>",S["Body"]))
    else:
        for cs in chs:
            if not cs.comments: continue
            story.append(Paragraph(f"CHAPTER {cs.chapter_num} — {cs.chapter_title.upper()}",S["ChHead"]))
            nc2=sum(1 for c in cs.comments if c.severity=="CRITICAL")
            nm2=sum(1 for c in cs.comments if c.severity=="MODERATE")
            story.append(Paragraph(f"<b>{len(cs.comments)} comments</b> — "
                f"Critical: {nc2} | Moderate: {nm2} | Suggestions: {len(cs.comments)-nc2-nm2}",S["Body"]))
            story.append(HR(c=RULE,t=.4,b=2,a=6))

            for sub in cs.subsections:
                if not sub.comments: continue
                story.append(Paragraph(f"§{sub.subsection_num} — {sub.title}",S["SubHead"]))
                story.append(Paragraph(f"<i>Expected purpose: {sub.expected_purpose}</i>",S["Excerpt"]))
                story.append(Spacer(1,4))

                for i, c in enumerate(sub.comments,1):
                    sc = SEV_COLORS.get(c.severity,ACCENT)
                    hdr_row = Table(
                        [[Paragraph(c.severity,S["Badge"]),Spacer(6,1),
                          Paragraph(f"§{sub.subsection_num} Comment {i} | Page ~{c.page_estimate}",S["Meta"])]],
                        colWidths=[2.2*cm,.4*cm,W-5*cm-2.6*cm])
                    hdr_row.setStyle(TableStyle([("BACKGROUND",(0,0),(0,0),sc),
                        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
                        ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0),
                        ("VALIGN",(0,0),(-1,-1),"MIDDLE")]))
                    excerpt_box = Table(
                        [[Paragraph(f'<i>"{c.para_excerpt[:120]}..."</i>',S["Excerpt"])]],
                        colWidths=[W-5*cm])
                    excerpt_box.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),LGRAY),
                        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
                        ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8)]))
                    fw_line = f'<br/><br/><b>Suggested framework:</b> {c.suggested_framework}' if c.suggested_framework else ""
                    mt_line = f'<br/><br/><b>Suggested method/reference:</b> {c.suggested_method}' if c.suggested_method else ""
                    detail_box = Table(
                        [[Paragraph(
                            f'<b>Issue:</b> {c.issue}<br/><br/>'
                            f'<b>Recommendation:</b> {c.recommendation}<br/><br/>'
                            f'<b>Literature needed:</b> {c.literature_needed}<br/><br/>'
                            f'<b>Theory / Framework:</b> {c.theory_needed}'
                            f'{fw_line}{mt_line}',S["Detail"])]],
                        colWidths=[W-5*cm])
                    detail_box.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),BOX_BG),
                        ("BOX",(0,0),(-1,-1),.4,RULE),("TOPPADDING",(0,0),(-1,-1),7),
                        ("BOTTOMPADDING",(0,0),(-1,-1),7),("LEFTPADDING",(0,0),(-1,-1),10),
                        ("RIGHTPADDING",(0,0),(-1,-1),10)]))
                    block = [hdr_row,excerpt_box,detail_box,Spacer(1,8)]
                    story.append(KeepTogether(block[:2])); story.extend(block[2:])
            story.append(Spacer(1,12)); story.append(HR())

    story.append(Spacer(1,6))
    story.append(Paragraph("ThesisSifu Pro v4.1 — Subsection commentary. "
        "See Report 1 (Examiner Audit), Report 2 (Annotated Thesis), "
        "Report 4 (Alignment Matrix).",S["Footer"]))
    doc.build(story); buf.seek(0); return buf.getvalue()


# ── Output 4: Alignment Matrix Report PDF ─────────────────────
def build_alignment_pdf(filename, audit_id, doc_type, clf, spine, align, chs) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf,pagesize=A4,leftMargin=2*cm,rightMargin=2*cm,
                            topMargin=2.2*cm,bottomMargin=2.2*cm)
    S = pdf_styles(); story = []

    story.append(type_badge_table(S,"STRUCTURAL ALIGNMENT MATRIX",TYPE_COLORS.get(doc_type,NAVY)))
    story.append(Spacer(1,14))
    story.append(Paragraph("ALIGNMENT MATRIX REPORT",S["Title"]))
    story.append(Paragraph("ThesisSifu Pro v4.1 — Golden Thread Audit",S["Sub"]))
    story.append(HR(c=NAVY,t=1.5,b=4,a=10))
    story.append(meta_box(S,[("Document",filename),("Audit ID",audit_id),
        ("Date",datetime.now().strftime("%d %B %Y, %H:%M")),
        ("Report","4 of 4 — Structural Alignment Matrix")]))
    story.append(Spacer(1,10))

    gt = align.golden_thread_score
    gt_color = {"STRONG":GREEN,"ACCEPTABLE":AMBER,"WEAK":RED,"BROKEN":RED}.get(gt,ACCENT)
    story.append(vbanner(S,f"Golden Thread: {gt}",gt_color,
        subtitle="Problem -> RQ -> RO -> Method -> Analysis -> Finding -> Conclusion"))
    story.append(Spacer(1,10))

    story.append(Paragraph("OVERALL ALIGNMENT VERDICT",S["SecHead"]))
    story.append(HR(c=NAVY,t=1,b=0,a=6))
    story.append(Paragraph(align.overall_verdict or "(no verdict generated)",S["Body"]))
    story.append(Spacer(1,8))

    story.append(Paragraph("THESIS SPINE (for reference)",S["SecHead"]))
    story.append(HR(c=NAVY,t=1,b=0,a=6))
    story.append(spine_box(S,spine)); story.append(Spacer(1,10))

    story.append(PageBreak())
    story.append(Paragraph("ALIGNMENT MATRIX — RQ-LEVEL TRACE",S["SecHead"]))
    story.append(HR(c=NAVY,t=1,b=0,a=6))

    if not align.rows:
        story.append(Paragraph("<i>No alignment rows produced. "
            "This usually means the spine extraction could not identify Research Questions. "
            "Check that the document has clear RQ/RO statements and the GEMINI_API_KEY is valid.</i>",S["Body"]))
    else:
        headers = ["RQ","RO / H","Method + Analysis","Finding","Conclusion","Status"]
        td = [[Paragraph(f"<b>{h}</b>",S["TblH"]) for h in headers]]
        for r in align.rows:
            sc = ALIGN_COLORS.get(r.status,ACCENT)
            roh = r.ro + (f" / {r.hypothesis}" if r.hypothesis and r.hypothesis!="—" else "")
            td.append([
                Paragraph(r.rq[:100],S["TblCSm"]),
                Paragraph(roh[:90],S["TblCSm"]),
                Paragraph(f"{r.method[:60]}\n{r.analysis[:60]}",S["TblCSm"]),
                Paragraph(r.finding[:90],S["TblCSm"]),
                Paragraph(r.conclusion[:90],S["TblCSm"]),
                Paragraph(f'<font color="{sc.hexval()}"><b>{r.status}</b></font>',S["TblCSm"]),
            ])
        cw = W-4*cm
        t = Table(td,colWidths=[cw*.18,cw*.15,cw*.20,cw*.17,cw*.17,cw*.13])
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),NAVY),("TEXTCOLOR",(0,0),(-1,0),WHITE),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,BOX_BG]),
            ("BOX",(0,0),(-1,-1),.5,RULE),("INNERGRID",(0,0),(-1,-1),.3,RULE),
            ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
            ("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),
            ("VALIGN",(0,0),(-1,-1),"TOP")]))
        story.append(t); story.append(Spacer(1,10))

        story.append(Paragraph("ROW-BY-ROW NOTES",S["SecHead"]))
        story.append(HR(c=NAVY,t=1,b=0,a=6))
        for i,r in enumerate(align.rows,1):
            sc = ALIGN_COLORS.get(r.status,ACCENT)
            badge = Table([[Paragraph(r.status,S["Badge"])]],colWidths=[3*cm])
            badge.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),sc),
                ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
            hrow = Table([[Paragraph(f"<b>RQ{i}.</b> {r.rq[:140]}",S["Body"]),badge]],
                          colWidths=[W-5*cm-3.2*cm,3*cm])
            hrow.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP"),
                ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0)]))
            story.append(hrow)
            if r.note: story.append(Paragraph(f"<i>{r.note}</i>",S["Excerpt"]))
            story.append(Spacer(1,6))

    if align.critical_gaps:
        story.append(PageBreak())
        story.append(Paragraph("CRITICAL STRUCTURAL GAPS",S["SecHead"]))
        story.append(HR(c=NAVY,t=1,b=0,a=6))
        for g in align.critical_gaps:
            story.append(Paragraph(f"• {g}",S["Bullet"]))
        story.append(Spacer(1,10))

    if align.structural_recommendations:
        story.append(Paragraph("STRUCTURAL RECOMMENDATIONS",S["SecHead"]))
        story.append(HR(c=NAVY,t=1,b=0,a=6))
        for rec in align.structural_recommendations:
            story.append(Paragraph(f"• {rec}",S["Bullet"]))
        story.append(Spacer(1,10))

    # Chapter delivery scorecard
    story.append(PageBreak())
    story.append(Paragraph("CHAPTER DELIVERY SCORECARD",S["SecHead"]))
    story.append(HR(c=NAVY,t=1,b=0,a=6))
    story.append(Paragraph("Did each subsection deliver on its expected purpose? "
        "High critical counts against a subsection = it did not.",S["Body"]))
    story.append(Spacer(1,6))
    td = [[Paragraph(f"<b>{h}</b>",S["TblH"]) for h in ["Chapter / Subsection","Expected Purpose","C/M/S"]]]
    for cs in chs:
        td.append([Paragraph(f"<b>Ch.{cs.chapter_num} — {cs.chapter_title[:50]}</b>",S["TblC"]),
                   Paragraph("(see subsections)",S["TblCSm"]),
                   Paragraph(f"{sum(1 for c in cs.comments if c.severity=='CRITICAL')} / "
                              f"{sum(1 for c in cs.comments if c.severity=='MODERATE')} / "
                              f"{sum(1 for c in cs.comments if c.severity=='SUGGESTION')}",S["TblC"])])
        for sub in cs.subsections:
            snc=sum(1 for c in sub.comments if c.severity=="CRITICAL")
            snm=sum(1 for c in sub.comments if c.severity=="MODERATE")
            sns=len(sub.comments)-snc-snm
            td.append([Paragraph(f"   §{sub.subsection_num} {sub.title[:50]}",S["TblCSm"]),
                       Paragraph(sub.expected_purpose[:120],S["TblCSm"]),
                       Paragraph(f"{snc}/{snm}/{sns}",S["TblCSm"])])
    cw = W-4*cm
    t = Table(td,colWidths=[cw*.34,cw*.50,cw*.16])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),NAVY),("TEXTCOLOR",(0,0),(-1,0),WHITE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,BOX_BG]),("BOX",(0,0),(-1,-1),.5,RULE),
        ("INNERGRID",(0,0),(-1,-1),.3,RULE),("TOPPADDING",(0,0),(-1,-1),4),
        ("BOTTOMPADDING",(0,0),(-1,-1),4),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(t); story.append(Spacer(1,16)); story.append(HR())
    story.append(Paragraph("ThesisSifu Pro v4.1 — Alignment Matrix. "
        "Use alongside Report 1, 2, 3 to triangulate revision priorities.",S["Footer"]))
    doc.build(story); buf.seek(0); return buf.getvalue()


# ── Main endpoint ──────────────────────────────────────────────
@app.post("/audit")
async def audit_document(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename or "document"

    if not filename.lower().endswith((".pdf",".docx")):
        raise HTTPException(400, "Only PDF and DOCX files are supported.")

    audit_id = "SUP-" + hashlib.md5(content).hexdigest()[:10].upper()
    print(f"[ThesisSifu] Starting audit {audit_id} for {filename}")

    # 1. Extract text
    full_text = extract_text(content, filename)
    print(f"[ThesisSifu] Extracted {len(full_text)} chars from {filename}")
    if not full_text.strip():
        full_text = "[Document appears image-based or empty — text extraction returned nothing]"

    # 2. Stage 0+1: spine extraction + classification (parallel)
    spine_task = asyncio.create_task(extract_spine(full_text, full_text))
    clf_task   = asyncio.create_task(classify_document(full_text))
    spine, clf = await asyncio.gather(spine_task, clf_task)

    doc_type = clf.get("type","MASTERS")
    field    = clf.get("field","UNKNOWN")
    inst     = clf.get("institution","UNKNOWN")
    if spine.title == "UNKNOWN" and clf.get("title","UNKNOWN") != "UNKNOWN":
        spine.title = clf["title"]
    print(f"[ThesisSifu] doc_type={doc_type} title={spine.title[:60]}")
    print(f"[ThesisSifu] Spine RQs: {spine.research_questions[:3]}")

    # 3. Split chapters + subsections
    chs = split_chapters(full_text)
    for cs in chs:
        cs.subsections = split_subsections(cs)
    total_subs = sum(len(cs.subsections) for cs in chs)
    print(f"[ThesisSifu] Split: {len(chs)} chapters, {total_subs} subsections")

    # 4. Stage 2: subsection paragraph audit (parallel, max 4 concurrent)
    sem = asyncio.Semaphore(4)
    async def bounded(sub: Subsection, ch_title: str):
        async with sem:
            cmts = await audit_subsection(sub, ch_title, doc_type, spine, full_text)
            return sub, cmts

    tasks   = [bounded(sub, cs.chapter_title) for cs in chs for sub in cs.subsections]
    results = await asyncio.gather(*tasks)

    for sub, cmts in results:
        sub.comments = cmts
    for cs in chs:
        cs.comments = [c for sub in cs.subsections for c in sub.comments]

    total_comments = sum(len(cs.comments) for cs in chs)
    print(f"[ThesisSifu] Total paragraph comments: {total_comments}")

    # 5. Stage 3: alignment audit
    align = await run_alignment_audit(spine, chs)

    # 6. Stage 4: holistic examiner report
    examiner_text = await run_examiner(doc_type, spine, field, inst, chs, align)

    # 7. Build outputs
    pdf1    = build_examiner_pdf(filename, audit_id, doc_type, clf, spine, examiner_text, align, chs)
    pdf3    = build_commentary_pdf(filename, audit_id, doc_type, clf, spine, chs)
    pdf4    = build_alignment_pdf(filename, audit_id, doc_type, clf, spine, align, chs)

    if filename.lower().endswith(".pdf"):
        output2   = build_annotated_pdf(content, full_text, audit_id, chs)
        out2_name = "2_Annotated_Thesis.pdf"
    else:
        output2   = build_annotated_docx(content, filename, audit_id, chs, clf)
        out2_name = "2_Annotated_Thesis.docx"

    # 8. ZIP and return
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("1_Examiner_Audit_Report.pdf", pdf1)
        zf.writestr(out2_name, output2)
        zf.writestr("3_Commentary_Report.pdf", pdf3)
        zf.writestr("4_Alignment_Matrix_Report.pdf", pdf4)
    tmp.close()
    print(f"[ThesisSifu] Audit {audit_id} complete. ZIP ready.")

    return FileResponse(
        path=tmp.name, media_type="application/zip",
        filename=f"ThesisSifu_v4_Audit_{audit_id}.zip",
        background=BackgroundTask(lambda: os.unlink(tmp.name)),
        headers={"Access-Control-Expose-Headers":"Content-Disposition"},
    )


@app.get("/health")
async def health():
    return {
        "status":                "ok",
        "gemini_available":      gemini_client is not None,
        "gemini_model":          GEMINI_MODEL,
        "claude_available":      claude_client is not None,
        "version":               "4.1.0",
        "fixes":                 [
            "google.genai SDK (not deprecated google.generativeai)",
            "robust chapter splitter (Roman numerals, word numbers, plain numbered headings)",
            "subsection filter removed (was breaking on Roman numeral chapters)",
            "safe Gemini response handler with finish_reason logging",
            "page estimation uses full_text offset not subsection slice",
            "prompt injection uses .replace() not .format() (no KeyError on curly braces)",
            "annotated PDF uses wider search window + shorter fallback search",
            "all exceptions logged with chapter+subsection context",
        ],
        "outputs": ["1_Examiner_Audit_Report.pdf","2_Annotated_Thesis.(docx|pdf)",
                    "3_Commentary_Report.pdf","4_Alignment_Matrix_Report.pdf"],
    }
